# affidavit_writer.py

from docx import Document
from datetime import datetime

def generate_affidavit(doc_a, doc_b, score, reasons, output_path):
    document = Document()

    document.add_heading("Forensic Affidavit of PDF Similarity", 0)

    document.add_paragraph(f"Prepared by: David Garzotto")
    document.add_paragraph(f"Title: Founder, Forensix, LLC")
    document.add_paragraph("Date: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    document.add_paragraph("Certification: Certified by Mile2 Investigations")
    document.add_paragraph("Statement: This analysis was conducted using best practices in digital forensics and PDF document examination, consistent with the evidentiary methodology referenced in Swigdoc and similar frameworks.")

    document.add_heading("Documents Compared", level=1)

    table = document.add_table(rows=3, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = doc_a["filename"]
    hdr_cells[2].text = doc_b["filename"]

    def row(label, a_val, b_val):
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = a_val or "N/A"
        row_cells[2].text = b_val or "N/A"

    row("SHA-256", doc_a["sha256"], doc_b["sha256"])
    row("XMP DocumentID", doc_a["xmp_document_id"], doc_b["xmp_document_id"])
    row("XMP InstanceID", doc_a["xmp_instance_id"], doc_b["xmp_instance_id"])
    row("Creation Date", doc_a["creation_date"], doc_b["creation_date"])
    row("Modification Date", doc_a["mod_date"], doc_b["mod_date"])

    document.add_heading("AcroForm Fields", level=2)
    if doc_a["form_fields"] or doc_b["form_fields"]:
        for d in [("A", doc_a), ("B", doc_b)]:
            document.add_paragraph(f"Form fields in Document {d[0]}:", style="List Bullet")
            for f in d[1]["form_fields"]:
                document.add_paragraph(f"{f['name']}: {f['value']}", style="List Number")
    else:
        document.add_paragraph("No form fields found.")

    document.add_heading("Similarity Score & Findings", level=1)
    document.add_paragraph(f"Match Score: {score}")
    if reasons:
        for r in reasons:
            document.add_paragraph(r, style="List Bullet")
    else:
        document.add_paragraph("No major similarities found.")

    document.add_paragraph("\nSignature: __________________________")
    document.add_paragraph("David Garzotto, Founder, Forensix, LLC")

    document.save(output_path)